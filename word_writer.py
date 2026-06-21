"""
RTL Hebrew Word document generator — full rewrite.

RTL fixes applied at 4 layers:
  1. Document settings  : <w:bidi/> in settings.xml  (document-wide RTL)
  2. Normal style       : <w:bidi/> on Normal pPr    (all paragraphs inherit RTL)
  3. Each paragraph     : <w:bidi/> + jc=right + pPr rPr lang he-IL
  4. Each run           : <w:rtl/> + lang he-IL + cs font + csSize

Bullet fix: Unicode • with w:ind right=360/hanging=360 (no ListBullet style).
Table RTL : <w:bidiVisual/> — col[0] is rightmost visually.
"""

import os, re
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.opc.constants import RELATIONSHIP_TYPE as RT

# ── Constants ─────────────────────────────────────────────────────────────────
FONT        = "Arial"
BODY_PT     = 12
H1_PT       = 16
H2_PT       = 14
H3_PT       = 12
TABLE_PT    = 11
SMALL_PT    = 9

C_H1            = "1F497D"   # dark blue
C_H2            = "2E75B6"   # medium blue
C_BODY          = "000000"
C_WHITE         = "FFFFFF"
C_GRAY          = "595959"
C_HEADER_BG     = "1F497D"   # table header background
C_ROW_ALT       = "EDF2FF"   # alternating table row
C_CALLOUT_BG    = "E8F4FD"   # recommendation box background
C_CALLOUT_BORD  = "2E75B6"   # recommendation box right border
C_TITLE_LINE    = "1F497D"   # title bottom border

OUTPUT_DIR  = os.path.join(os.path.dirname(__file__), "output")
OUTPUT_FILE = os.path.join(OUTPUT_DIR, "המלצות_הדרכות_AI.docx")

# ── XML micro-helpers ─────────────────────────────────────────────────────────

def _find_or_create(parent, tag):
    el = parent.find(qn(tag))
    if el is None:
        el = OxmlElement(tag)
        parent.append(el)
    return el


def _set_color(rPr, hex_color):
    _find_or_create(rPr, "w:color").set(qn("w:val"), hex_color)


def _set_shd(pPr_or_tcPr, fill_hex):
    shd = _find_or_create(pPr_or_tcPr, "w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  fill_hex)


# ── Core RTL primitives ───────────────────────────────────────────────────────

def _para_rtl(para):
    """Apply RTL at the paragraph level (4 things)."""
    pPr = para._p.get_or_add_pPr()

    # 1. <w:bidi/>
    if pPr.find(qn("w:bidi")) is None:
        pPr.insert(0, OxmlElement("w:bidi"))

    # 2. right-align
    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # 3. paragraph-level default run language = Hebrew
    pRPr = _find_or_create(pPr, "w:rPr")
    _find_or_create(pRPr, "w:lang").set(qn("w:bidi"), "he-IL")


def _run_rtl(run, pt, bold=False, color=C_BODY):
    """Apply RTL + Hebrew + Arial to a run."""
    rPr = run._r.get_or_add_rPr()

    # Fonts: latin + complex-script (Hebrew)
    rFonts = _find_or_create(rPr, "w:rFonts")
    for attr in ("w:ascii", "w:hAnsi", "w:cs"):
        rFonts.set(qn(attr), FONT)

    # Size: both latin and complex-script (half-points)
    hp = str(int(pt * 2))
    _find_or_create(rPr, "w:sz"  ).set(qn("w:val"), hp)
    _find_or_create(rPr, "w:szCs").set(qn("w:val"), hp)

    # Bold
    if bold:
        _find_or_create(rPr, "w:b")
        _find_or_create(rPr, "w:bCs")

    # Color
    _set_color(rPr, color)

    # RTL run direction
    _find_or_create(rPr, "w:rtl")

    # Language = Hebrew
    _find_or_create(rPr, "w:lang").set(qn("w:bidi"), "he-IL")


# ── Document setup ────────────────────────────────────────────────────────────

def _setup_doc(doc):
    """
    Set RTL at EVERY level of the document:
      Layer 1 — settings.xml          : <w:bidi/> + themeFontLang he-IL
      Layer 2 — styles.xml docDefaults: pPrDefault gets bidi + jc=right
                                        rPrDefault lang fixed to he-IL (was ar-SA)
      Layer 3 — Normal style           : bidi + jc=right
      Layer 4 — each paragraph/run    : applied in _para_rtl / _run_rtl
    """
    # ── Layer 1: settings.xml ────────────────────────────────────────────────
    s = doc.settings.element
    if s.find(qn("w:bidi")) is None:
        s.insert(0, OxmlElement("w:bidi"))
    _find_or_create(s, "w:themeFontLang").set(qn("w:bidi"), "he-IL")

    # ── Layer 2: docDefaults in styles.xml ───────────────────────────────────
    styles_el   = doc.styles.element
    doc_def     = styles_el.find(qn("w:docDefaults"))
    if doc_def is not None:
        # --- paragraph defaults: add bidi + right-align ---
        pPr_def = doc_def.find(qn("w:pPrDefault"))
        if pPr_def is None:
            pPr_def = OxmlElement("w:pPrDefault")
            doc_def.append(pPr_def)
        pPr = pPr_def.find(qn("w:pPr"))
        if pPr is None:
            pPr = OxmlElement("w:pPr")
            pPr_def.append(pPr)
        if pPr.find(qn("w:bidi")) is None:
            pPr.insert(0, OxmlElement("w:bidi"))
        _find_or_create(pPr, "w:jc").set(qn("w:val"), "right")

        # --- run defaults: fix language from ar-SA to he-IL ---
        rPr_def = doc_def.find(qn("w:rPrDefault"))
        if rPr_def is not None:
            rPr = rPr_def.find(qn("w:rPr"))
            if rPr is not None:
                lang = rPr.find(qn("w:lang"))
                if lang is not None:
                    lang.set(qn("w:val"),  "he-IL")
                    lang.set(qn("w:bidi"), "he-IL")

    # ── Layer 3: Normal style ────────────────────────────────────────────────
    norm_pPr = doc.styles["Normal"].element.get_or_add_pPr()
    if norm_pPr.find(qn("w:bidi")) is None:
        norm_pPr.insert(0, OxmlElement("w:bidi"))
    _find_or_create(norm_pPr, "w:jc").set(qn("w:val"), "right")

    # ── A4 page ──────────────────────────────────────────────────────────────
    sec = doc.sections[0]
    sec.page_width    = Cm(21)
    sec.page_height   = Cm(29.7)
    sec.left_margin   = Cm(2.5)
    sec.right_margin  = Cm(2.5)
    sec.top_margin    = Cm(2.5)
    sec.bottom_margin = Cm(2.5)


# ── Visual paragraph builders ─────────────────────────────────────────────────

def _clean(t):
    t = re.sub(r"\*\*(.*?)\*\*", r"\1", t)
    t = re.sub(r"\*(.*?)\*",     r"\1", t)
    return t.strip()


def add_title_block(doc, title, subtitle=None):
    """Large RTL title with a blue bottom border line."""
    para = doc.add_paragraph()
    _para_rtl(para)

    # Bottom border (thick blue line under the title)
    pPr = para._p.get_or_add_pPr()
    pBdr = _find_or_create(pPr, "w:pBdr")
    bot = _find_or_create(pBdr, "w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "18")
    bot.set(qn("w:color"), C_TITLE_LINE)
    bot.set(qn("w:space"), "4")

    run = para.add_run(_clean(title))
    _run_rtl(run, 22, bold=True, color=C_H1)
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after  = Pt(4)

    if subtitle:
        s = doc.add_paragraph()
        _para_rtl(s)
        sr = s.add_run(_clean(subtitle))
        _run_rtl(sr, BODY_PT, color=C_GRAY)
        s.paragraph_format.space_after = Pt(16)


def add_divider(doc):
    """Thin gray horizontal rule."""
    para = doc.add_paragraph()
    pPr  = para._p.get_or_add_pPr()
    pBdr = _find_or_create(pPr, "w:pBdr")
    bot  = _find_or_create(pBdr, "w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "4")
    bot.set(qn("w:color"), "CCCCCC")
    bot.set(qn("w:space"), "2")
    para.paragraph_format.space_before = Pt(8)
    para.paragraph_format.space_after  = Pt(8)


def add_tool_heading(doc, number, name):
    """Colored section header: dark-blue background, white text."""
    para = doc.add_paragraph()
    _para_rtl(para)

    pPr = para._p.get_or_add_pPr()
    _set_shd(pPr, C_HEADER_BG)

    # Indentation inside the colored box
    ind = _find_or_create(pPr, "w:ind")
    ind.set(qn("w:right"), "200")
    ind.set(qn("w:left"),  "200")

    run = para.add_run(f"{number}. {name}")
    _run_rtl(run, H2_PT, bold=True, color=C_WHITE)
    para.paragraph_format.space_before = Pt(18)
    para.paragraph_format.space_after  = Pt(6)


def add_h1(doc, text):
    para = doc.add_paragraph()
    _para_rtl(para)
    run = para.add_run(_clean(text))
    _run_rtl(run, H1_PT, bold=True, color=C_H1)
    para.paragraph_format.space_before = Pt(20)
    para.paragraph_format.space_after  = Pt(6)


def add_h2(doc, text):
    para = doc.add_paragraph()
    _para_rtl(para)
    run = para.add_run(_clean(text))
    _run_rtl(run, H2_PT, bold=True, color=C_H2)
    para.paragraph_format.space_before = Pt(14)
    para.paragraph_format.space_after  = Pt(4)


def add_h3(doc, text):
    para = doc.add_paragraph()
    _para_rtl(para)
    run = para.add_run(_clean(text))
    _run_rtl(run, H3_PT, bold=True, color=C_BODY)
    para.paragraph_format.space_before = Pt(10)
    para.paragraph_format.space_after  = Pt(2)


def add_body(doc, text, italic=False):
    text = _clean(text)
    if not text:
        return
    para = doc.add_paragraph()
    _para_rtl(para)
    run = para.add_run(text)
    _run_rtl(run, BODY_PT, color=C_BODY)
    if italic:
        run.font.italic = True
    para.paragraph_format.space_after = Pt(4)


def add_bullet(doc, text):
    """RTL bullet: • + hanging indent from the right."""
    text = _clean(text)
    if not text:
        return
    para = doc.add_paragraph()
    _para_rtl(para)

    pPr = para._p.get_or_add_pPr()
    ind = _find_or_create(pPr, "w:ind")
    ind.set(qn("w:right"),   "360")
    ind.set(qn("w:hanging"), "360")

    run = para.add_run("•  " + text)
    _run_rtl(run, BODY_PT)
    para.paragraph_format.space_after = Pt(3)


def add_callout(doc, title, text):
    """
    Highlighted recommendation box:
    light-blue background + thick blue right border.
    """
    para = doc.add_paragraph()
    _para_rtl(para)

    pPr = para._p.get_or_add_pPr()

    # Background
    _set_shd(pPr, C_CALLOUT_BG)

    # Thick right border (= start of RTL text)
    pBdr = _find_or_create(pPr, "w:pBdr")
    right = _find_or_create(pBdr, "w:right")
    right.set(qn("w:val"),   "single")
    right.set(qn("w:sz"),    "24")
    right.set(qn("w:color"), C_CALLOUT_BORD)
    right.set(qn("w:space"), "6")

    ind = _find_or_create(pPr, "w:ind")
    ind.set(qn("w:right"), "160")
    ind.set(qn("w:left"),  "160")

    r1 = para.add_run(title + ": ")
    _run_rtl(r1, BODY_PT, bold=True, color=C_H1)

    r2 = para.add_run(text)
    _run_rtl(r2, BODY_PT, color=C_BODY)

    para.paragraph_format.space_before = Pt(8)
    para.paragraph_format.space_after  = Pt(14)


# ── RTL Table ─────────────────────────────────────────────────────────────────

def _add_hyperlink(paragraph, url, display_text, pt=SMALL_PT, color="2E75B6"):
    """
    Inserts a real clickable hyperlink into a paragraph.
    Automatically prepends https:// if the url has no scheme.
    """
    if not url.startswith(("http://", "https://")):
        url = "https://" + url

    # Register external relationship → returns rId
    r_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)

    # <w:hyperlink r:id="rIdX">
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    # <w:r> inside the hyperlink
    run_el = OxmlElement("w:r")
    rPr    = OxmlElement("w:rPr")

    # Font
    rFonts = OxmlElement("w:rFonts")
    for attr in ("w:ascii", "w:hAnsi", "w:cs"):
        rFonts.set(qn(attr), FONT)
    rPr.append(rFonts)

    # Size
    hp = str(int(pt * 2))
    for tag in ("w:sz", "w:szCs"):
        el = OxmlElement(tag)
        el.set(qn("w:val"), hp)
        rPr.append(el)

    # Underline
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)

    # Color
    clr = OxmlElement("w:color")
    clr.set(qn("w:val"), color)
    rPr.append(clr)

    # RTL + Hebrew language
    rPr.append(OxmlElement("w:rtl"))
    lang = OxmlElement("w:lang")
    lang.set(qn("w:bidi"), "he-IL")
    rPr.append(lang)

    run_el.append(rPr)

    # Text
    t = OxmlElement("w:t")
    t.text = display_text
    run_el.append(t)

    hyperlink.append(run_el)
    paragraph._p.append(hyperlink)


def _cell_rtl(cell, text, pt=TABLE_PT, bold=False, color=C_BODY, bg=None, small_text=None, rating=None):
    """Write RTL text into a table cell (first paragraph)."""
    p = cell.paragraphs[0]
    _para_rtl(p)

    if bg:
        tcPr = cell._tc.get_or_add_tcPr()
        _set_shd(tcPr, bg)
        # Also set cell vertical alignment to center
        vAlign = _find_or_create(tcPr, "w:vAlign")
        vAlign.set(qn("w:val"), "center")

    # Add top/bottom cell padding
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = _find_or_create(tcPr, "w:tcMar")
    for side in ("w:top", "w:bottom"):
        m = _find_or_create(tcMar, side)
        m.set(qn("w:w"),    "80")
        m.set(qn("w:type"), "dxa")
    for side in ("w:right", "w:left"):
        m = _find_or_create(tcMar, side)
        m.set(qn("w:w"),    "120")
        m.set(qn("w:type"), "dxa")

    run = p.add_run(text)
    _run_rtl(run, pt, bold=bold, color=color)

    if small_text:
        p2 = cell.add_paragraph()
        _para_rtl(p2)
        _add_hyperlink(p2, small_text, small_text, pt=SMALL_PT, color="2E75B6")
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after  = Pt(0)

    if rating:
        p3 = cell.add_paragraph()
        _para_rtl(p3)
        r3 = p3.add_run(rating)
        _run_rtl(r3, SMALL_PT, color="595959")
        p3.paragraph_format.space_before = Pt(2)
        p3.paragraph_format.space_after  = Pt(0)


def add_courses_table(doc, courses):
    """
    RTL table: col[0]=rightmost (course name), col[1]=description,
                col[2]=level, col[3]=duration.
    <w:bidiVisual/> reverses visual column order.
    Headers: dark blue bg, white bold.
    Rows: alternating white / very-light-blue.
    """
    headers = ["שם הקורס", "מה תלמד", "רמה", "משך"]

    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"

    # ── Table-level RTL ──────────────────────────────────────────────────────
    tbl  = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    if tbl.tblPr is None:
        tbl.insert(0, tblPr)

    # bidiVisual → columns render right-to-left
    if tblPr.find(qn("w:bidiVisual")) is None:
        tblPr.append(OxmlElement("w:bidiVisual"))

    # Table width = full text area
    tblW = _find_or_create(tblPr, "w:tblW")
    tblW.set(qn("w:w"),    "9072")   # twips (~16 cm)
    tblW.set(qn("w:type"), "dxa")

    # Table borders (light gray grid)
    tblBorders = _find_or_create(tblPr, "w:tblBorders")
    for side in ("w:top","w:left","w:bottom","w:right","w:insideH","w:insideV"):
        b = _find_or_create(tblBorders, side)
        b.set(qn("w:val"),   "single")
        b.set(qn("w:sz"),    "4")
        b.set(qn("w:color"), "CCCCCC")

    # ── Column widths (twips, right→left: name, desc, level, dur) ────────────
    widths = [2520, 4320, 1116, 1116]   # total = 9072 twips ≈ 16 cm
    tblGrid = OxmlElement("w:tblGrid")
    for w in widths:
        gc = OxmlElement("w:gridCol")
        gc.set(qn("w:w"), str(w))
        tblGrid.append(gc)
    tbl.insert(1, tblGrid)   # insert after tblPr

    # ── Header row ────────────────────────────────────────────────────────────
    hrow = table.rows[0]
    for cell, hdr in zip(hrow.cells, headers):
        _cell_rtl(cell, hdr, pt=TABLE_PT, bold=True, color=C_WHITE, bg=C_HEADER_BG)

    # ── Data rows ─────────────────────────────────────────────────────────────
    for i, c in enumerate(courses):
        row = table.add_row()
        bg  = C_ROW_ALT if i % 2 == 0 else "FFFFFF"
        _cell_rtl(row.cells[0], c["name"], pt=TABLE_PT, bold=True, bg=bg,
                  small_text=c.get("url"))
        _cell_rtl(row.cells[1], c["desc"], pt=TABLE_PT, bg=bg, rating=c.get("rating"))
        _cell_rtl(row.cells[2], c["level"], pt=TABLE_PT, bg=bg)
        _cell_rtl(row.cells[3], c["dur"],  pt=TABLE_PT, bg=bg)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)


# ── Markdown fallback parser (used by crew output) ────────────────────────────

def _parse_markdown(doc, content):
    for line in content.splitlines():
        s = line.strip()
        if not s:
            continue
        if   s.startswith("# "):    add_h1(doc, s[2:])
        elif s.startswith("## "):   add_h2(doc, s[3:])
        elif s.startswith("### "):  add_h3(doc, s[4:])
        elif s.startswith(("- ","* ","• ")):
            add_bullet(doc, s[2:])
        elif re.match(r"^\d+\.\s", s):
            add_bullet(doc, re.sub(r"^\d+\.\s+", "", s))
        else:
            add_body(doc, s)


# ── Public API ────────────────────────────────────────────────────────────────

def create_word_document(content: str, output_path: str = None) -> str:
    """Markdown → RTL Hebrew Word (used by crew output)."""
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    out = output_path or OUTPUT_FILE
    doc = Document()
    _setup_doc(doc)
    _parse_markdown(doc, content)
    doc.save(out)
    return out


def new_document() -> Document:
    """Return a fresh RTL-configured Document (used by generate_report.py)."""
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    doc = Document()
    _setup_doc(doc)
    return doc
